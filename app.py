from flask import Flask, render_template, request, redirect, url_for, flash
from werkzeug.utils import secure_filename
from docx import Document
from unidecode import unidecode
import os

app = Flask(__name__)
app.secret_key = "supersecretkey"
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Разрешенные расширения файлов
ALLOWED_EXTENSIONS = {'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text(filename):
    doc = Document(filename)
    text = ''.join(para.text for para in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += ''.join(para.text for para in cell.paragraphs)
    return text

def calculate_time(total_chars):
    total_seconds = total_chars * (150 / 1800)
    return round(total_seconds / 60)

def format_time(minutes):
    hours = minutes // 60
    mins = minutes % 60
    return f"{hours:02d}:{mins:02d}"

@app.route('/', methods=['GET', 'POST'])
def index():
    # Инициализация данных формы
    form_data = {
        'cost_dictator': '',
        'cost_studio': '',
        'cost_sound_engineer': '',
        'cost_extra_services': '',
        'discount': '',
        'manual_chars': '',
        'current_chars': '',           # Сохранённое количество символов
        'prev_filename': '',           # Имя загруженного файла (для отображения)
        'prev_storage_filename': '',   # Безопасное имя файла (для хранения)
        'use_manual_chars': False      # Режим ручного ввода
    }
    result = None

    if request.method == 'POST':
        # Сохраняем данные формы (если новое значение для manual_chars или current_chars не пришло, оставляем старое)
        form_data['cost_dictator'] = request.form.get('cost_dictator', '')
        form_data['cost_studio'] = request.form.get('cost_studio', '')
        form_data['cost_sound_engineer'] = request.form.get('cost_sound_engineer', '')
        form_data['cost_extra_services'] = request.form.get('cost_extra_services', '')
        form_data['discount'] = request.form.get('discount', '')
        form_data['manual_chars'] = request.form.get('manual_chars', form_data['manual_chars'])
        form_data['current_chars'] = request.form.get('current_chars', form_data['current_chars'])

        try:
            cost_dictator = float(form_data['cost_dictator'])
        except ValueError:
            cost_dictator = 0.0
        try:
            cost_studio = float(form_data['cost_studio'])
        except ValueError:
            cost_studio = 0.0
        try:
            cost_sound_engineer = float(form_data['cost_sound_engineer'])
        except ValueError:
            cost_sound_engineer = 0.0
        try:
            cost_extra_services = float(form_data['cost_extra_services']) if form_data['cost_extra_services'].strip() != '' else 0.0
        except ValueError:
            cost_extra_services = 0.0
        try:
            discount = float(form_data['discount']) if form_data['discount'].strip() != '' else None
        except ValueError:
            discount = None

        manual_chars_str = form_data['manual_chars']
        if manual_chars_str.strip() != '':
            try:
                manual_chars = int(manual_chars_str)
                form_data['use_manual_chars'] = True
            except ValueError:
                manual_chars = None
                form_data['use_manual_chars'] = False
        else:
            manual_chars = None

        # Обработка загруженного файла
        file = request.files.get('file')
        file_uploaded = False
        filepath = None
        if file and file.filename:
            # Проверка расширения файла
            if not allowed_file(file.filename):
                flash("Пожалуйста, загрузите файл в формате .docx")
                return redirect(url_for('index'))
            original_filename = file.filename
            filename = secure_filename(unidecode(original_filename))
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            form_data['prev_filename'] = original_filename
            form_data['prev_storage_filename'] = filename
            file_uploaded = True
            form_data['use_manual_chars'] = False  # Новый файл имеет приоритет
            form_data['manual_chars'] = ""         # Сброс ручного ввода
        else:
            hidden_storage = request.form.get('prev_storage_filename')
            if hidden_storage:
                filename_storage = hidden_storage
                filename_display = request.form.get('prev_filename', '')
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename_storage)
                form_data['prev_filename'] = filename_display
                form_data['prev_storage_filename'] = filename_storage

        # Определяем источник количества символов:
        if file_uploaded:
            try:
                text = extract_text(filepath)
            except Exception as e:
                flash(f"Ошибка при чтении файла: {e}")
                return redirect(url_for('index'))
            computed_chars = len(text)
            total_chars = computed_chars
            form_data['current_chars'] = str(computed_chars)
            manual_mode = False
        elif form_data['use_manual_chars'] and manual_chars is not None:
            total_chars = manual_chars
            form_data['current_chars'] = str(manual_chars)
            manual_mode = True
        elif form_data['current_chars'].strip() != '':
            total_chars = int(form_data['current_chars'])
            manual_mode = form_data['use_manual_chars']
        else:
            flash("Пожалуйста, загрузите файл или введите количество символов вручную")
            return redirect(url_for('index'))

        total_minutes = calculate_time(total_chars)
        recording_time = max(total_minutes * 2, 60)
        processing_time = max(total_minutes * 6, 120)

        recording_time_hours = recording_time / 60
        processing_time_hours = processing_time / 60
        cost_dictator_total = cost_dictator * recording_time_hours
        cost_studio_total = cost_studio * recording_time_hours
        cost_sound_engineer_total = (cost_sound_engineer * recording_time_hours) + (cost_sound_engineer * processing_time_hours)
        total_cost = cost_studio_total + cost_dictator_total + cost_sound_engineer_total + cost_extra_services

        result = {
            'total_chars': total_chars,
            'total_minutes': format_time(total_minutes),
            'recording_time': format_time(recording_time),
            'processing_time': format_time(processing_time),
            'cost_dictator': round(cost_dictator_total, 2),
            'cost_studio': round(cost_studio_total, 2),
            'cost_sound_engineer': round(cost_sound_engineer_total, 2),
            'total_cost': round(total_cost, 2)
        }
        if not manual_mode:
            result['filename'] = form_data['prev_filename']
        if cost_extra_services:
            result['cost_extra_services'] = round(cost_extra_services, 2)
        if discount is not None and discount > 0:
            discounted_cost = total_cost * (1 - discount / 100)
            result['discounted_cost'] = round(discounted_cost, 2)

    return render_template('index.html', result=result, form_data=form_data)

if __name__ == '__main__':
    app.run(debug=True)
